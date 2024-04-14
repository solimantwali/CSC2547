import csv
import json
import re
import textwrap
from abc import ABC
from pathlib import Path
from pprint import pprint

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src import constants

patterns=[
    "GI",
    "Persuade",
    "Persuade with",
    "Q",
    "SR",
    "CR",
    "AF",
    "Seek",
    "Emphasize",
    "Confront"
]

class Conversation(ABC):
    pass


class SimpleConversation(Conversation):
    def __init__(
        self,
        counsellor_bot,
        client_bot,
        eval_bot,
        end_classifier,
        output_dir,
        mode,
        price_limit=3,
        turn_limit=100,
    ):
        self.counsellor_bot = counsellor_bot
        self.client_bot = client_bot
        self.eval_bot = eval_bot
        self.end_classifier = end_classifier
        self.chat_log = []
        self.output_dir = output_dir
        self.mode = mode
        self.state = "counsellor"
        self.counsellor_end = False
        self.client_end = False
        self.prompt_tokens = dict()
        self.completion_tokens = dict()
        self.price = 0
        self.price_limit = price_limit
        self.total_turns = 0
        self.turn_limit = turn_limit

    def _after_each_turn(self, turn):
        self.chat_log.append(turn)
        pprint(turn)
        match self.mode:
            case "normal":
                observer_response = input("Continue? [y]/n: ")
                if "n" in observer_response.lower():
                    self._save_and_exit()
            case "turbo":
                pass

    def _save_as_docx(self):
        doc = Document()
        doc.add_heading("Counsellor System Prompt", level=1)
        paragraph = doc.add_paragraph(self.counsellor_bot.system_prompt)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_heading("Client System Prompt", level=1)
        paragraph = doc.add_paragraph(self.client_bot.system_prompt)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if self.eval_bot:
            doc.add_heading("Evaluator System Prompt", level=1)
            paragraph = doc.add_paragraph(self.eval_bot.system_prompt)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_heading("Conversation", level=1)
        for turn in self.chat_log:
            paragraph = doc.add_paragraph()
            interlocutor = paragraph.add_run(turn["name"].title() + ": ")
            interlocutor.bold = True
            paragraph.add_run(turn["content"])
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        doc.save(self.output_dir / "conversation.docx")

    def _save_as_txt(self):
        wrapper = textwrap.TextWrapper(width=88)
        with open(self.output_dir / "conversation.txt", "w", encoding="utf-8") as file:
            file.write("Counsellor System Prompt:\n")
            file.write(self.counsellor_bot.system_prompt)
            file.write("\n" + "*" * 88 + "\n")
            file.write("\n\nClient System Prompt:\n")
            file.write(self.client_bot.system_prompt)
            file.write("\n" + "*" * 88 + "\n")
            if self.eval_bot:
                file.write("\n\nEvaluator System Prompt:\n")
                file.write(self.eval_bot.system_prompt)
                file.write("\n" + "*" * 88 + "\n")
            file.write("\n\nConversation:\n\n")
            for turn in self.chat_log:
                file.write(wrapper.fill(turn["name"].title() + ": " + turn["content"]))
                file.write("\n\n")

    def _save_and_exit(self):
        self.report_turns()
        self.report_price()
        with open(self.output_dir / "counsellor.json", "w", encoding="utf-8") as file:
            json.dump(self.counsellor_bot.conversation_history, file, indent=4)
        with open(self.output_dir / "client.json", "w", encoding="utf-8") as file:
            json.dump(self.client_bot.conversation_history, file, indent=4)
        if len(self.counsellor_bot.to_csv) > 0:
            with open(
                self.output_dir / "counsellor_bc_history.csv", "w", newline=""
            ) as file:
                writer = csv.writer(file)
                writer.writerow(["Utterance", "Counsellor_BC"])
                writer.writerows(self.counsellor_bot.to_csv)
        if self.eval_bot:
            with open(
                self.output_dir / "evaluator.json", "w", encoding="utf-8"
            ) as file:
                json.dump(self.eval_bot.conversation_history, file, indent=4)
            with open(
                self.output_dir / "evaluator_bc_history.csv", "w", newline=""
            ) as file:
                writer = csv.writer(file)
                writer.writerow(["Utterance", "Evaluator_BC"])
                writer.writerows(self.eval_bot.to_csv)
        self._save_as_docx()
        self._save_as_txt()
        exit()

    def _save_and_exit_eval(self):
        self.report_price()
        match self.eval_bot.version:
            case constants.EVALUATOR_CSV_VER:
                fn = "evaluator_csv.json"
            case constants.EVALUATOR_GLOBAL_SCORE_VER:
                fn = "evaluator_global_score.json"
            case _:
                fn = "evaluator.json"

        with open(self.output_dir / fn, "w", encoding="utf-8") as file:
            json.dump(self.eval_bot.conversation_history, file, indent=4)
        exit()

    def _print_prompts(self):
        prompts = {
            "client_system_prompt": self.client_bot.system_prompt,
            "counsellor_system_prompt": self.counsellor_bot.system_prompt,
            "evaluator_system_prompt": self.eval_bot.system_prompt
            if self.eval_bot
            else None,
        }
        pprint(prompts)
        with open(Path(self.output_dir) / "prompts.json", "w") as file:
            json.dump(prompts, file, indent=4)

    def add_price(self, usage, model):
        # Computes the estimated price charged by OpenAI based on tokens used so far
        pricing_dict = {
            "gpt-4-0125-preview": (0.01, 0.03),
            "gpt-4-1106-preview": (0.01, 0.03),
            "gpt-4-1106-vision-preview": (0.01, 0.03),
            "gpt-4": (0.03, 0.06),
            "gpt-4-32k": (0.06, 0.12),
            "gpt-3.5-turbo-0125": (0.0005, 0.0015),
            "gpt-3.5-turbo-instruct": (0.0015, 0.0020),
            "gpt-3.5-turbo-1106": (0.0010, 0.0020),
            "gpt-3.5-turbo-0613": (0.0015, 0.0020),
            "gpt-3.5-turbo-16k-0613": (0.0030, 0.0040),
            "gpt-3.5-turbo-0301": (0.0015, 0.0020),
        }

        if model in self.prompt_tokens:
            self.prompt_tokens[model] += usage.prompt_tokens
        else:
            self.prompt_tokens[model] = usage.prompt_tokens

        if model in self.completion_tokens:
            self.completion_tokens[model] += usage.completion_tokens
        else:
            self.completion_tokens[model] = usage.completion_tokens

        if model in pricing_dict:
            input_price, output_price = pricing_dict[model]
        else:
            print(f"ERROR: MODEL UNKNOWN ({model})")
            self._save_and_exit()

        self.price += (input_price / 1000 * usage.prompt_tokens) + (
            output_price / 1000 * usage.completion_tokens
        )

    def report_price(self):
        print("Total tokens:")
        for model in self.prompt_tokens:
            print(f"Model: {model}")
            print(f"    Input tokens: {self.prompt_tokens[model]}")
            print(f"    Output tokens: {self.completion_tokens[model]}")
        print(f"Price charged: {self.price}")

    def report_turns(self):
        print(f"Total turns (not including eval): {self.total_turns}")

    def bc_count(self,eval_hist):
        # print(eval_hist)
        data = eval_hist
        # with open(eval_hist) as file:
        #     data = json.load(file)

        aggr_metrics = {}
        aggr_metrics["name"] = "aggregated metrics"
        aggr_metrics["codes"] = {p: 0 for p in patterns}
        aggr_metrics["summary metrics"] = {}
        

        turn = 1
        for d in data[1:]:
            if d["name"] != "evaluator":
                # data.remove(d)
                continue
            multi_encode = {p: 0 for p in patterns}
            for i, pattern in enumerate(patterns):
                if re.search(pattern, d["content"]):
                    multi_encode[pattern] = 1
            aggr_metrics["codes"] = {
                x: aggr_metrics["codes"].get(x, 0) + multi_encode.get(x, 0)
                for x in patterns
            }

        aggr_metrics["summary metrics"]["R:Q"] = (aggr_metrics["codes"]["CR"] + aggr_metrics["codes"]["SR"]) / aggr_metrics["codes"]["Q"] if aggr_metrics["codes"]["Q"] > 0 else None
        aggr_metrics["summary metrics"]["%CR"] = (
            (aggr_metrics["codes"]["CR"]) / (aggr_metrics["codes"]["CR"] + aggr_metrics["codes"]["SR"]) * 100 if aggr_metrics["codes"]["CR"] > 0 else None
        )
        return aggr_metrics

    def amendments(aggr_metrics):
        amend = []
        if aggr_metrics["summary metrics"]["%CR"] < 0.4:
            amend.append("You are using too many SR and not enough CR. If you choose to output a reflection, make it complex (CR)")
        if aggr_metrics["summary metrics"]["R:Q"] < 1:
            amend.append("You are asking too many questions and not reflecting enough. Consider SR or CR instead of Q.")
        if aggr_metrics["summary metrics"]["Persuade"] > 0:
            amend.append("You used Persuade, a MI Non-adherent Code. Refrain from repeating similar sentiments")
        if aggr_metrics["summary metrics"]["Confront"] > 0:
            amend.append("You used Confront, a MI Non-adherent Code. Refrain from repeating similar sentiments")

    def converse(
        self,
        client_openai_obj,
        counsellor_openai_obj,
        evaluator_openai_obj,
        end_classifier_openai_obj,
    ):
        counsellor_turn = None
        client_turn = None
        running_counts = {}
        self._print_prompts()

        try:
            while self.state != "end":
                match self.state:
                    case "counsellor":
                        print(f"counsellor's turn now. evalbot record is:", self.bc_count(self.eval_bot.conversation_history))
                        counsellor_response, usage = self.counsellor_bot.converse(
                            openai_obj=counsellor_openai_obj, client_turn=client_turn
                        )
                        counsellor_turn = {
                            "role": "user",
                            "name": "counsellor",
                            "content": counsellor_response,
                        }
                        self.total_turns += 1
                        self._after_each_turn(counsellor_turn)

                        self.add_price(usage, self.counsellor_bot.model)

                        if end_classifier_openai_obj:
                            # Classify counsellor's intent of ending
                            end_cls_response, usage = self.end_classifier.converse(
                                openai_obj=end_classifier_openai_obj,
                                input=counsellor_turn,
                            )
                            if "true" in end_cls_response.lower():
                                self.counsellor_end = True
                            elif self.client_end:
                                # if client wants to end but counsellor doesn't,
                                # then reset flags
                                # TODO: should counsellor always respect client?
                                self.client_end = False

                            self.add_price(usage, self.end_classifier.model)

                        if self.price >= self.price_limit:
                            next_state = "end"
                            print("Price limit reached")
                            self._save_and_exit()
                        elif self.total_turns >= self.turn_limit:
                            next_state = "end"
                            print("Turn limit reached")
                            self._save_and_exit()
                        elif self.counsellor_end and self.client_end:
                            next_state = "end"
                            print("Conversation ended naturally")
                            self._save_and_exit()
                        elif evaluator_openai_obj:
                            next_state = "evaluator"
                        else:
                            next_state = "client"

                    case "client":
                        client_response, usage = self.client_bot.converse(
                            openai_obj=client_openai_obj,
                            counsellor_turn=counsellor_turn,
                        )
                        client_turn = {
                            "role": "user",
                            "name": "client",
                            "content": client_response,
                        }
                        self.total_turns += 1
                        self._after_each_turn(client_turn)

                        self.add_price(usage, self.client_bot.model)

                        if end_classifier_openai_obj:
                            # Classify counsellor's intent of ending
                            end_cls_response, usage = self.end_classifier.converse(
                                openai_obj=end_classifier_openai_obj, input=client_turn
                            )
                            if "true" in end_cls_response.lower():
                                self.client_end = True
                            elif self.counsellor_end:
                                # if counsellor wants to end but client doesn't,
                                # then reset flags
                                self.counsellor_end = False

                            self.add_price(usage, self.end_classifier.model)

                        if self.price >= self.price_limit:
                            next_state = "end"
                            print("Price limit reached")
                            self._save_and_exit()
                        elif self.total_turns >= self.turn_limit:
                            next_state = "end"
                            print("Turn limit reached")
                            self._save_and_exit()
                        elif self.counsellor_end and self.client_end:
                            next_state = "end"
                            print("Conversation ended naturally")
                            self._save_and_exit()
                        else:
                            next_state = "counsellor"

                    case "evaluator":

                        evaluator_response, usage = self.eval_bot.converse(
                            openai_obj=evaluator_openai_obj,
                            counsellor_turn=counsellor_turn,
                            client_turn=client_turn,
                        )
                        evaluator_turn = {
                            "role": "user",
                            "name": "evaluator",
                            "content": evaluator_response,
                        }
                        self._after_each_turn(evaluator_turn)

                        self.add_price(usage, self.eval_bot.model)

                        if self.price >= self.turn_limit:
                            next_state = "end"
                            print("Price limit reached")
                            self._save_and_exit()
                        # Technically don't need to check anything
                        # other than price, but just in case
                        elif self.total_turns >= self.turn_limit:
                            next_state = "end"
                            print("Turn limit reached")
                            self._save_and_exit()
                        elif self.counsellor_end and self.client_end:
                            next_state = "end"
                            print("Conversation ended naturally")
                            self._save_and_exit()
                        else:
                            next_state = "client"

                self.state = next_state
        except KeyboardInterrupt:
            print("Conversation ended")
            self._save_and_exit()

    def evaluate(
        self,
        evaluator_openai_obj,
        transcript_json,
    ):
        transcript = json.load(transcript_json)
        idx = 1
        while idx < len(transcript):
            last_two_turns = []
            while transcript[idx]["name"] != "counsellor":
                last_two_turns.append(transcript[idx])
                idx += 1
                if idx >= len(transcript):
                    break

            if idx >= len(transcript):
                break
            last_two_turns.append(transcript[idx])

            print(last_two_turns)

            evaluator_response, usage = self.eval_bot.evaluate(
                openai_obj=evaluator_openai_obj, last_two_turns=last_two_turns
            )

            print(evaluator_response)

            evaluator_turn = {
                "role": "user",
                "name": "evaluator",
                "content": evaluator_response,
            }

            self.chat_log.extend(last_two_turns)
            self.chat_log.append(evaluator_turn)

            self.add_price(usage, self.eval_bot.model)
            if self.price >= self.turn_limit:
                print("Price limit reached")
                self._save_and_exit_eval()

            idx += 1

        print("Evaluation finished")
        self._save_and_exit_eval()

    def evaluate_csv(self, evaluator_openai_obj, transcript_csv):
        # pprint(self.eval_bot.system_prompt)
        reader = csv.reader(transcript_csv)
        next(reader)

        with open(
            self.output_dir / "result.csv", mode="w", newline="", encoding="utf-8"
        ) as file:
            writer = csv.writer(file)
            header = ["Utterance", "Output", "Label"]
            writer.writerow(header)

            counter = 0
            correct_counter = 0
            for row in reader:
                if len(row) > 2:
                    if row[2] != "TRUE" or row[4] != "FALSE":
                        continue

                utterance = row[0]
                # print(utterance)

                utterance_dict = {
                    "role": "user",
                    "name": "counsellor",
                    "content": utterance,
                }
                evaluator_response, usage = self.eval_bot.evaluate(
                    openai_obj=evaluator_openai_obj,
                    last_two_turns=[utterance_dict],
                    greedy=True,
                )

                # print(evaluator_response)

                evaluator_turn = {
                    "role": "user",
                    "name": "evaluator",
                    "content": evaluator_response,
                }

                try:
                    output_label = re.findall(r"\(([^)]+)\)", evaluator_response)[0]
                except IndexError:
                    output_label = evaluator_response.strip()

                label = row[1].strip()
                output_row = [utterance, output_label, label]

                counter += 1
                if label == output_label:
                    correct_counter += 1

                writer.writerow(output_row)

                self.chat_log.extend(row[0])
                self.chat_log.append(evaluator_turn)

                self.add_price(usage, self.eval_bot.model)
                if self.price >= self.turn_limit:
                    print("Price limit reached")
                    self._save_and_exit_eval()

        print("Evaluation finished")
        print("Accuracy: {}".format(format(correct_counter / counter, "0.2%")))
        self._save_and_exit_eval()

    def evaluate_globalscores(self, evaluator_openai_obj, transcript_json):
        pprint(self.eval_bot.system_prompt)
        transcript = json.load(transcript_json)
        evaluator_response, usage = self.eval_bot.evaluate(
            openai_obj=evaluator_openai_obj, last_two_turns=transcript
        )

        print(evaluator_response)

        evaluator_turn = {
            "role": "user",
            "name": "evaluator",
            "content": evaluator_response,
        }

        self.chat_log.extend(evaluator_response)
        self.chat_log.append(evaluator_turn)
        self.eval_bot.conversation_history = evaluator_turn

        self.add_price(usage, self.eval_bot.model)
        if self.price >= self.turn_limit:
            print("Price limit reached")
            self._save_and_exit_eval()

        print("Evaluation finished")
        self._save_and_exit_eval()
